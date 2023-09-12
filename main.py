import logging
import os
from flask import Flask, request, send_file
from flask_cors import CORS
from supabase_client import create_supabase_client
import whisper
from docx import Document
from moviepy.editor import VideoFileClip, TextClip, CompositeVideoClip
from moviepy.video.tools.subtitles import file_to_subtitles
from pycaption import SRTReader, SRTWriter
import zipfile
import datetime

logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
CORS(app)

model = whisper.load_model("base")


def seconds_to_srt_time(seconds):
    h, r = divmod(seconds, 3600)
    m, r = divmod(r, 60)
    s, ms = divmod(r, 1)
    ms = int(ms * 1000)
    return f"{int(h):02}:{int(m):02}:{int(s):02},{ms:03}"


def docx_to_srt(docx_filename, srt_filename):
    doc = Document(docx_filename)
    subtitles = []
    index = 1
    for paragraph in doc.paragraphs:
        if '-->' in paragraph.text:
            time, text = paragraph.text.split(']', 1)
            start_time, end_time = time[1:].split(' --> ')
            start_time = seconds_to_srt_time(float(start_time))
            end_time = seconds_to_srt_time(float(end_time))
            subtitles.append(f"{index}\n{start_time} --> {end_time}\n{text.strip()}\n\n")
            index += 1
    with open(srt_filename, 'w', encoding='utf-8-sig') as out_file:
        out_file.writelines(subtitles)


def write_srt_file(filename, captions):
    with open(filename, 'w', encoding='utf-8-sig') as f:
        f.write(SRTWriter().write(captions))


def read_srt_file(filename):
    with open(filename, 'r', encoding='utf-8-sig') as f:
        return SRTReader().read(f.read())


def read_srt_file_as_generator(filename):
    with open(filename, 'r', encoding='utf-8-sig') as f:
        return file_to_subtitles(f)


def create_text_clips(captions):
    for caption in captions.get_captions('en-US'):
        start = caption.start / 1000000
        end = caption.end / 1000000
        txt = caption.get_text()
        txt_clip = TextClip(txt, font='FreeMono', fontsize=24, color='white')
        txt_clip = txt_clip.set_duration(end - start)
        yield (start, txt_clip)


@app.route('/transcribe', methods=['POST'])
def download_file():
    bucket_name = request.json.get('bucketName')
    source = request.json.get('source')

    try:
        supabase = create_supabase_client()
    except Exception as e:
        logging.error("Error creating Supabase client: %s", e)
        return 'Error creating Supabase client', 500

    destination = os.path.join(os.getcwd(), os.path.basename(source))

    try:
        with open(destination, 'wb+') as f:
            res = supabase.storage.from_(bucket_name).download(source)
            f.write(res)
    except Exception as e:
        logging.error("Error downloading file: %s", e)
        return 'Error downloading file', 500

    docx_filename = "transcription.docx"
    srt_filename = "transcription.srt"

    doc = Document()
    result = model.transcribe(destination, fp16=False, verbose=True)

    for segment in result["segments"]:
        start_time = str(segment["start"])
        end_time = str(segment["end"])
        text = segment["text"]
        doc.add_paragraph(f"[{start_time} --> {end_time}]  {text}")

    doc.save(docx_filename)
    docx_to_srt(docx_filename, srt_filename)

    captions = read_srt_file(srt_filename)
    clips = list(create_text_clips(captions))

    clip = VideoFileClip(destination)
    final = CompositeVideoClip([clip] + [c[1].set_start(c[0]).set_position(('center', 'bottom')) for c in clips])

    output_filename = f"captioned_{os.path.basename(source)}"
    final.write_videofile(output_filename)

    transcription_filename = f"{os.path.splitext(os.path.basename(source))[0]}.txt"
    transcription_path = os.path.join(os.getcwd(), transcription_filename)

    with open(transcription_path, 'w', encoding='utf-8') as f:
        f.write(result["text"])

    user_id = os.path.splitext(os.path.basename(transcription_filename))[0].split('_')[0]

    current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"{user_id}_captioned_{current_time}.zip"
    with zipfile.ZipFile(zip_filename, 'w') as zip_file:
        zip_file.write(output_filename)
        zip_file.write(srt_filename)
        zip_file.write(transcription_filename)

    try:
        with open(zip_filename, 'rb') as f:
            supabase.storage.from_(bucket_name).upload(zip_filename, f)
    except Exception as e:
        logging.error("Error uploading zip file to Supabase: %s", e)
        return 'Error uploading zip file to Supabase', 500

    try:
        response = send_file(zip_filename, as_attachment=True)
        response.headers["Content-Disposition"] = f"attachment; filename=captioned_files.zip"
        return response
    except Exception as e:
        logging.error("Error sending zip file as a response: %s", e)
        return 'Error sending zip file as a response', 500


if __name__ == '__main__':
    app.run(port=5000)